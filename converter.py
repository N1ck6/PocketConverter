import os
import sys
from PIL import Image
from docx import Document
from io import StringIO
from moviepy import VideoFileClip
from win11toast import toast
import traceback

all_list = ['mp4', 'gif', 'txt', 'pdf', 'docx', 'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff', 'ico']
log_file_path = "C:\\Program Files\\PocketConverter\\PocketConverter_log.txt"
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)
icon_dir = os.path.join(base_path, 'logo.ico')
font_dir = os.path.join(base_path, 'DejaVuSansCondensed.ttf')


def create_name(filepath, mode):
    new_filepath = f"{filepath[0] + filepath[1]}.{mode}"
    if mode != 'pngs' and os.path.isfile(new_filepath):
        for i in range(1, 101):
            name = f"{filepath[1]}({i})"
            if not os.path.isfile(f"{filepath[0] + name}.{mode}"):
                return name
        toast("File name already exists", "Try renaming the file", icon=icon_dir)
        sys.exit()
    return filepath[1]


def convert_file(filepath, mode):
    if "folder" in mode:
        if not os.path.isdir(filepath):
            toast("Folder doesn't exist", "How and Why?", icon=icon_dir)
            sys.exit()
        return convert_document(filepath, mode)
    if not os.path.isfile("".join(filepath)):
        toast("File doesn't exist", "How and Why?", icon=icon_dir)
        sys.exit()
    ext = filepath[2]
    if ext == "." + mode:
        toast("File with that extension already exists", "How and Why?", icon=icon_dir)
        sys.exit()
    new_name = create_name(filepath, mode)
    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.ico']:
        convert_image(filepath, mode, new_name)
    elif ext in ['.txt', '.pdf', '.docx']:
        convert_document(filepath, mode, new_name)
    elif ext in ['.mp4', '.gif']:
        convert_animated(filepath, mode, new_name)
    else:
        toast("Format not allowed right now", "Maybe in next updates...", icon=icon_dir)
        sys.exit()
    return new_name


def convert_image(filepath, mode, new_name):
    img = Image.open("".join(filepath))
    if mode == "jpg" or 'A' not in img.getbands():
        img = img.convert("RGB")
    new_filepath = f"{filepath[0] + new_name}.{mode}"
    img.save(new_filepath)


def convert_document(filepath, mode, new_name=""):
    if "folder" in mode:
        if "pdf" in mode:
            return images_to_pdf(filepath)
        elif "gif" in mode:
            return images_to_gif(filepath)
    elif filepath[2] == ".txt":
        if mode == "pdf":
            convert_txt_to_pdf(filepath, new_name)
        elif mode == "docx":
            convert_txt_to_docx(filepath, new_name)
    elif filepath[2] == ".docx":
        if mode == "txt":
            convert_docx_to_txt(filepath, new_name)
        elif mode == "pdf":
            convert_docx_to_pdf(filepath, new_name)
    elif filepath[2] == ".pdf":
        if mode == "txt":
            convert_pdf_to_txt(filepath, new_name)
        elif mode == "docx":
            convert_pdf_to_docx(filepath, new_name)


def convert_txt_to_pdf(filepath, new_name):
    from fpdf import FPDF
    from textwrap import wrap
    pdf = FPDF()
    pdf.set_auto_page_break(True, margin=10)
    pdf.add_font('DejaVu', '', font_dir, uni=True)
    pdf.set_font('DejaVu', '', size=11)
    pdf.add_page()
    with open("".join(filepath), "r", encoding="utf-8") as f:
        text = f.read()
        for line in text.split('\n'):
            lines = wrap(line, 90)
            if len(lines) == 0:
                pdf.ln()
            for wrapped in lines:
                pdf.cell(0, 4.8, wrapped, ln=1, align="J")
    new_filepath = f"{filepath[0] + new_name}.pdf"
    pdf.output(new_filepath, 'F')


def convert_txt_to_docx(filepath, new_name):
    doc = Document()
    with open("".join(filepath), "r", encoding="utf-8") as f:
        doc.add_paragraph(f.read())
    new_filepath = f"{filepath[0] + new_name}.docx"
    doc.save(new_filepath)


def convert_docx_to_txt(filepath, new_name):
    new_filepath = f"{filepath[0] + new_name}.txt"
    doc = Document("".join(filepath))
    with open(new_filepath, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")


def convert_docx_to_pdf(filepath, new_name):
    from docx2pdf import convert
    new_filepath = f"{filepath[0] + new_name}.pdf"
    sys.stderr = StringIO()
    convert("".join(filepath), new_filepath)


def convert_pdf_to_txt(filepath, new_name):
    from pymupdf import open as pdfopen
    new_filepath = f"{filepath[0] + new_name}.txt"
    with open(new_filepath, "w", encoding="utf-8") as f:
        for page in pdfopen("".join(filepath)):
            f.write(page.get_text())


def convert_pdf_to_docx(filepath, new_name):
    from pdf2docx import Converter
    cv = Converter("".join(filepath))
    cv.convert(f"{filepath[0] + new_name}.docx", start=0, end=None)
    cv.close()


def images_to_pdf(folder):
    from fpdf import FPDF
    image_paths = [os.path.join(folder, i) for i in sorted([f for f in os.listdir(folder) if f.rsplit('.', 1)[1] in all_list[5:8]])]
    if not image_paths:
        toast("Folder doesn't have images", "Available format: png, jpg, jpeg", icon=icon_dir)
        sys.exit()
    pdf = FPDF()
    a4_w, a4_h = 210, 297
    for image_path in image_paths:
        pdf.add_page()
        width, height = Image.open(image_path).size
        scale = min(a4_w / width, a4_h / height)
        if scale < 1:
            width *= scale
            height *= scale
        pdf.image(image_path, (a4_w - width) / 2, (a4_h - height) / 2, width, height)
    filename = create_name([folder.rsplit('\\', 1)[0], 'Combined_images'], 'pdf')
    pdf.output(f"{folder.rsplit('\\', 1)[0]}\\{filename}.pdf")
    return filename


def images_to_gif(folder):
    filename = create_name([folder.rsplit('\\', 1)[0], 'Combined_images'], 'gif')
    output_path = f"{folder.rsplit('\\', 1)[0]}\\{filename}.gif"
    image_files = [f for f in os.listdir(folder) if f.rsplit('.', 1)[1] in all_list[5:10]]
    if not image_files:
        toast("Folder doesn't have images", "Available format: png, jpg, jpeg, bmp, webp", icon=icon_dir)
        sys.exit()
    images = [Image.open(os.path.join(folder, i)).convert('RGB') for i in sorted(image_files)]
    min_width = min(img.width for img in images)
    min_height = min(img.height for img in images)
    images = [img.resize((min_width, min_height)) for img in images]
    images[0].save(output_path, save_all=True, append_images=images[1:], duration=25 * len(images), loop=0)
    return filename


def convert_animated(filepath, mode, new_name):
    if filepath[2] == ".mp4":
        if mode == "gif":
            video_to_gif(filepath, new_name)
        elif mode == "mp3":
            video_to_audio(filepath, new_name)
    elif filepath[2] == ".gif":
        if mode == "pngs":
            gif_to_photos(filepath, new_name)
        elif mode == "png":
            gif_to_photo(filepath, new_name)
        elif mode == "mp4":
            gif_to_video(filepath, new_name)


def video_to_gif(filepath, new_name):
    new_filepath = f"{filepath[0] + new_name}.gif"
    clip = VideoFileClip("".join(filepath))
    clip.write_gif(new_filepath, logger=None)
    clip.close()


def video_to_audio(filepath, new_name):
    new_filepath = f"{filepath[0] + new_name}.mp3"
    video = VideoFileClip("".join(filepath))
    video.audio.write_audiofile(new_filepath, logger=None)
    video.close()


def gif_to_video(filepath, new_name):
    new_filepath = f"{filepath[0] + new_name}.mp4"
    video = VideoFileClip("".join(filepath))
    video.write_videofile(new_filepath, logger=None)
    video.close()


def gif_to_photo(filepath, new_name):
    frame = Image.open("".join(filepath))
    nframes = 0
    while frame:
        frame.save(f"{filepath[0] + new_name}.png", 'PNG')
        nframes += 1
        try:
            frame.seek(nframes)
            break
        except EOFError:
            break


def gif_to_photos(filepath, new_name):
    out_folder = os.path.join(filepath[0], "extracted gif")
    os.makedirs(out_folder, exist_ok=True)
    frame = Image.open("".join(filepath))
    nframes = 0
    while frame:
        frame.save(f"{out_folder}\\{new_name}-{nframes}.png", 'PNG')
        nframes += 1
        try:
            frame.seek(nframes)
        except EOFError:
            break


if __name__ == "__main__":
    try:
        devnull = open(os.devnull, 'w')
        sys.stdout, sys.stderr = devnull, devnull
        if len(sys.argv) == 3:
            file, mode = sys.argv[1], sys.argv[2]
            logs_input = f"Input! File: {file}, Mode: {mode}\n"
            if "folder" not in mode:
                file = [os.path.dirname(file) + '\\', os.path.basename(file).rsplit('.', 1)[0], os.path.splitext(file)[1]]
            name = convert_file(file, mode)
            if "folder" in mode:
                mode = mode[-3:]
            if mode == "pngs":
                toast(f"Done!", f"Saved as extracted gif folder", icon=icon_dir, group='done')
            else:
                toast(f"Done!", f"Saved as {name}.{mode}", icon=icon_dir, group='done')
        else:
            toast(f"Hi!", f"I need to be used through context menu ;)", icon=icon_dir, group='done')
    except Exception as e:
        if os.path.isfile(log_file_path) and logs_input and file and mode:
            with open(log_file_path, 'a') as log_file:  # Logs file to check what might have caused errors
                log_file.write(f"Error! Input: {logs_input} File: {file}, Mode: {mode}\n")
                log_file.write("Error Output:\n")
                log_file.write(traceback.format_exc())
                log_file.write('---------------------------------------------------------------------\n')
        toast("Conversion error", fr"Check log file in in C:\Program Files\PocketConverter for more information", icon=icon_dir)
