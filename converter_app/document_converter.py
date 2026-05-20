import re
from pathlib import Path
from converter_app.utils import show_toast, FONT_PATH
import xml.etree.ElementTree as ET

class DocumentConverter:
    SUPPORTED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md', 'markdown', 'html', 'rtf', 'odt', 'epub'}

    _UNICODE_FRACTIONS = {
        '½': '1/2', '¼': '1/4', '¾': '3/4',
        '⅓': '1/3', '⅔': '2/3', '⅕': '1/5',
        '⅖': '2/5', '⅗': '3/5', '⅘': '4/5',
        '⅙': '1/6', '⅚': '5/6', '⅛': '1/8',
        '⅜': '3/8', '⅝': '5/8', '⅞': '7/8',
        '⅑': '1/9', '⅒': '1/10', '⅟': '1/',
    }
    
    _RE_MARKDOWN_CLEAN = re.compile(
        r'(?m)^#{1,6}\s+|'  # Headings
        r'(?<!\\)\*\*(.+?)\*\*|'  # Bold **
        r'(?<!\\)__(.+?)__|'  # Bold __
        r'(?<!\\)\*(.+?)\*|'  # Italic *
        r'(?<!\\)_(.+?)_|'  # Italic _
        r'`(.+?)`|'  # Inline code
        r'\[(.+?)\]\(.+?\)|'  # Links [text](url)
        r'^>\s+|'  # Quote markers
        r'^[-*]\s+|'  # List markers
        r'~~(.+?)~~'  # Strikethrough
    )
    _RE_SQRT_N = re.compile(r'\\sqrt\[(\d+)\]\{([^}]+)\}')
    _RE_SQRT   = re.compile(r'\\sqrt\{([^}]+)\}')
    _RE_FRAC   = re.compile(r'\\frac\{([^}]+)\}\{([^}]+)\}')
    _RE_HTML_TAGS = re.compile(r'<[^>]+>')
    _RE_WHITESPACE = re.compile(r'\n{3,}')

    def convert(self, filepath: str, mode: str, new_name: str) -> None:
        ext = Path(filepath).suffix[1:].lower()
        mode = mode.lower()

        handlers = {
            'txt': self._convert_from_txt, 'docx': self._convert_from_docx,
            'pdf': self._convert_from_pdf, 'md': self._convert_from_markdown,
            'markdown': self._convert_from_markdown, 'html': self._convert_from_html
        }
        handler = handlers.get(ext)
        if not handler:
            show_toast("Error", f"Source format .{ext} not supported")
            return
        handler(filepath, mode, new_name)

    def _convert_from_txt(self, filepath: str, mode: str, new_name: str) -> None:
        processed = self._process_txt_with_gpt_support(filepath)
        out = Path(filepath).parent / f"{new_name}"
        if mode == 'pdf': self._write_pdf(out, processed)
        elif mode == 'docx': self._write_docx(out, processed)
        elif mode in ('md', 'markdown'): (out.with_suffix('.md')).write_text(processed, encoding='utf-8')
        elif mode == 'html': self._write_html(out, processed)
        else: show_toast("Error", f"Target format .{mode} not supported")

    def _convert_from_docx(self, filepath: str, mode: str, new_name: str) -> None:
        from docx import Document
        
        doc = Document(filepath)
        text = '\n'.join(p.text for p in doc.paragraphs)
        out = Path(filepath).parent / new_name

        if mode == 'txt': out.with_suffix('.txt').write_text(text, encoding='utf-8')
        elif mode == 'pdf': self._write_pdf(out, text)
        elif mode in ('md', 'markdown'): out.with_suffix('.md').write_text(text, encoding='utf-8')
        else: show_toast("Error", f"Target format .{mode} not supported")

    def _convert_from_pdf(self, filepath: str, mode: str, new_name: str) -> None:
        from pymupdf import open as pdfopen

        doc = pdfopen(filepath)
        text = ''.join(page.get_text('text') for page in doc)
        doc.close()
        out = Path(filepath).parent / new_name
        
        if mode == 'txt': out.with_suffix('.txt').write_text(text, encoding='utf-8')
        elif mode == 'docx': self._pdf_to_docx(filepath, new_name)
        elif mode == 'md': out.with_suffix('.md').write_text(text, encoding='utf-8')
        else: show_toast("Error", f"Target format .{mode} not supported")

    def _pdf_to_docx(self, filepath: str, new_name: str) -> None:
        from pdf2docx import Converter
        
        out = Path(filepath).parent / f"{new_name}.docx"
        cv = Converter(filepath)
        try: cv.convert(str(out), start=0, end=None)
        finally: cv.close()

    def _convert_from_markdown(self, filepath: str, mode: str, new_name: str) -> None:
        import markdown

        md_text = Path(filepath).read_text(encoding='utf-8')
        out = Path(filepath).parent / new_name
        
        if mode == 'pdf': self._write_pdf(out, self._html_to_text(markdown.markdown(md_text)))
        elif mode == 'html': (out.with_suffix('.html')).write_text(self._wrap_html(markdown.markdown(md_text, extensions=['extra', 'codehilite']), Path(filepath).stem), encoding='utf-8')
        elif mode == 'docx': self._markdown_to_docx(md_text, out)
        elif mode == 'txt': out.with_suffix('.txt').write_text(self._clean_markdown(md_text), encoding='utf-8')
        else: show_toast("Error", f"Target format .{mode} not supported")

    def _markdown_to_docx(self, text: str, out: Path) -> None:
        from docx import Document

        doc = Document()
        for line in text.split('\n'):
            line = line.rstrip()
            if line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('## '): doc.add_heading(line[3:], level=2)
            elif line.startswith('### '): doc.add_heading(line[4:], level=3)
            elif line.strip(): doc.add_paragraph(line)
            
        self._save_docx(doc, out)

    def _convert_from_html(self, filepath: str, mode: str, new_name: str) -> None:
        from bs4 import BeautifulSoup

        html = Path(filepath).read_text(encoding='utf-8')
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n')
        out = Path(filepath).parent / f"{new_name}"

        if mode == 'pdf': self._write_pdf(out, text)
        elif mode == 'txt': out.with_suffix('.txt').write_text(text, encoding='utf-8')
        elif mode == 'md': import html2text; out.with_suffix('.md').write_text(html2text.html2text(html), encoding='utf-8')
        else: show_toast("Error", f"Target format .{mode} not supported")

    def _write_pdf(self, out: Path, text: str) -> None:
        from fpdf import FPDF; from textwrap import wrap

        pdf = FPDF()
        pdf.set_auto_page_break(True, margin=10)
        pdf.add_font('DejaVu', '', str(FONT_PATH), uni=True)
        pdf.set_font('DejaVu', '', size=11)
        pdf.add_page()

        for line in text.split('\n'):
            for wrapped in wrap(line, 90) or ['']: pdf.cell(0, 4.8, wrapped, ln=1)
        pdf.output(str(out.with_suffix('.pdf')))

    def _write_docx(self, out: Path, text: str) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph(text)
        self._save_docx(doc, out)
        
    def _repackage_clean_docx(self, src_path: Path, dst_path: Path) -> None: # Repackage docx, removing thumbnail/customXml bloat
        import zipfile

        ESSENTIAL_FILES = {
            '[Content_Types].xml',
            '_rels/.rels',
            'word/document.xml',
            'word/_rels/document.xml.rels',
            'word/styles.xml',
            'word/settings.xml',
            'word/webSettings.xml',
            'word/fontTable.xml',
            'word/theme/theme1.xml',
            'docProps/core.xml',
            'docProps/app.xml',
        }

        with zipfile.ZipFile(src_path, 'r') as src_z:
            with zipfile.ZipFile(dst_path, 'w', zipfile.ZIP_DEFLATED) as dst_z:
                for item in src_z.infolist():
                    if item.filename in ESSENTIAL_FILES:
                        data = src_z.read(item.filename)

                        # Fix _rels/.rels: remove thumbnail relationship
                        if item.filename == '_rels/.rels':
                            data = self._remove_thumbnail_rels(data)

                        # Fix [Content_Types].xml: remove customXml/thumbnail entries
                        if item.filename == '[Content_Types].xml':
                            data = self._clean_content_types(data)

                        # Fix word/_rels/document.xml.rels: remove customXml/numbering/stylesWithEffects
                        if item.filename == 'word/_rels/document.xml.rels':
                            data = self._clean_document_rels(data)

                        dst_z.writestr(item, data)
                        
    def _remove_thumbnail_rels(self, rels_xml: bytes) -> bytes: 
        root = ET.fromstring(rels_xml)
        ns = {'': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/relationships')

        for rel in root.findall('Relationship', ns):
            rel_type = rel.get('Type', '')
            if 'thumbnail' in rel_type.lower():
                root.remove(rel)

        return ET.tostring(root, encoding='UTF-8', xml_declaration=True)

    def _clean_content_types(self, ct_xml: bytes) -> bytes:
        root = ET.fromstring(ct_xml)
        ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/content-types')

        for elem in list(root):
            part_name = elem.get('PartName', '')
            ext = elem.get('Extension', '')
            if ('customXml' in part_name or 
                'thumbnail' in part_name.lower() or
                'stylesWithEffects' in part_name or
                'numbering' in part_name):
                root.remove(elem)
            if ext.lower() == 'jpeg':
                root.remove(elem)

        return ET.tostring(root, encoding='UTF-8', xml_declaration=True)

    def _clean_document_rels(self, rels_xml: bytes) -> bytes:
        root = ET.fromstring(rels_xml)
        ns = {'': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/relationships')

        for rel in list(root.findall('Relationship', ns)):
            rel_type = rel.get('Type', '').lower()
            target = rel.get('Target', '').lower()
            if ('customxml' in rel_type or
                'numbering' in rel_type or
                'styleswitheffects' in rel_type or
                'customxml' in target or
                'numbering' in target or
                'styleswitheffects' in target):
                root.remove(rel)

        return ET.tostring(root, encoding='UTF-8', xml_declaration=True)

    def _save_docx(self, doc, out: Path) -> None:
        temp_path = out.with_suffix('.tmp.docx')
        doc.save(str(temp_path))

        clean_path = out.with_suffix('.docx')
        self._repackage_clean_docx(temp_path, clean_path)

        if temp_path.exists():
            temp_path.unlink()
    
    def _write_html(self, out: Path, text: str) -> None:
        (out.with_suffix('.html')).write_text(
            f"<!DOCTYPE html>\n<html>\n<head><meta charset='utf-8'></head>\n<body>\n{text}\n</body>\n</html>",
            encoding='utf-8')

    def _html_to_text(self, html: str) -> str:
        return self._RE_HTML_TAGS.sub('\n', html).replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')

    def _wrap_html(self, body: str, title: str) -> str:
        return f"""<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title>
               <style>body{{font-family:sans-serif;max-width:800px;margin:0 auto;padding:20px}}
               code{{background:#f4f4f4;padding:2px 5px;border-radius:3px}}
               pre{{background:#f4f4f4;padding:10px;border-radius:5px;overflow-x:auto}}</style></head>
               <body>{body}</body></html>"""

    def _process_txt_with_gpt_support(self, filepath: str) -> str:
        try:
            if Path(filepath).exists():
                content = Path(filepath).read_text(encoding='utf-8')
            else:
                content = filepath
        except (OSError, ValueError):
            content = filepath
        cleaned = self._clean_markdown(content)
        cleaned = self._RE_WHITESPACE.sub('\n\n', cleaned)
        return self._convert_formulas(cleaned).strip()
    
    def _clean_markdown(self, text: str) -> str:
        """Remove markdown syntax while preserving content."""
        def _replace(match):
            if match.lastindex:
                for i in range(1, match.lastindex + 1):
                    group = match.group(i)
                    if group is not None:
                        return group
            return ''
        return self._RE_MARKDOWN_CLEAN.sub(_replace, text)

    def _convert_formulas(self, text: str) -> str:
        for unicode_frac, ascii_frac in self._UNICODE_FRACTIONS.items(): # unicode fractions
            text = text.replace(unicode_frac, ascii_frac)
        text = re.sub(r'³√\(([^)]+)\)', r'(\1)^(1/3)', text) # cube root
        text = re.sub(r'√\(([^)]+)\)', r'(\1)^(1/2)', text) # square root with parentheses
        text = re.sub(r'√(\d+)', r'(\1)^(1/2)', text) # square root without parentheses
        text = self._RE_SQRT_N.sub(r'(\2)^(1/\1)', text) # sqrt with index
        text = self._RE_SQRT.sub(r'(\1)^(1/2)', text) # sqrt
        text = self._RE_FRAC.sub(r'(\1)/(\2)', text) # fractions
        text = (text.replace('×', '*').replace('÷', '/') # math symbols
                .replace('−', '-').replace('±', '+/-')
                .replace('π', 'pi').replace('∞', 'inf')
                .replace('\\left(', '(').replace('\\right)', ')'))
        text = re.sub(r'\\(sin|cos|tan|log|ln|exp)', r'\1', text)

        return text

    @staticmethod
    def merge_pdfs(pdf_files: list[str], output_path: str) -> None:
        from pypdf import PdfMerger
        
        merger = PdfMerger()
        for f in pdf_files:
            merger.append(f)
        merger.write(output_path)
        merger.close()

    @staticmethod
    def split_pdf(pdf_path: str, output_folder: str, pages_per_file: int = 1) -> None:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(pdf_path)
        out_dir = Path(output_folder)
        out_dir.mkdir(exist_ok=True)
        total = len(reader.pages)

        for i in range(0, total, pages_per_file):
            writer = PdfWriter()
            for j in range(i, min(i + pages_per_file, total)): writer.add_page(reader.pages[j])
            with open(out_dir / f"split_{i//pages_per_file + 1}.pdf", "wb") as f: writer.write(f)